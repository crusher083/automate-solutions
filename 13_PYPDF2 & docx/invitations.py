from docx import Document
from docx.enum.text import WD_BREAK

with open ('guests.txt', 'r') as f:
    names = [line.rstrip() for line in f]
doc = Document()
for n in names:
    doc.add_paragraph(f'{n}', 'Title')
    doc.add_paragraph('Invitation to connference!', 'Quote')
    doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)
doc.save('invitations.docx')